#!/usr/bin/env python3
"""
Document Processing Module
Supports: PDF, DOCX, TXT, MD
"""

from pathlib import Path
from typing import Dict, List, Optional
import os


def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text = ""

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            text += f"\n--- Page {i+1} ---\n{page_text}\n"

        return text.strip()
    except ImportError:
        return "Error: pypdf not installed. Run: pip install pypdf"
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document

        doc = Document(file_path)
        text = ""

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text += para.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"

        return text.strip()
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except UnicodeDecodeError:
        # Try different encodings
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read().strip()
        except Exception as e:
            return f"Error reading TXT: {str(e)}"
    except Exception as e:
        return f"Error reading TXT: {str(e)}"


def extract_markdown_text(file_path: str) -> str:
    """Extract text from Markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Optionally convert markdown to plain text
        try:
            import markdown
            html = markdown.markdown(content)
            # Strip HTML tags for plain text
            import re
            text = re.sub('<[^<]+?>', '', html)
            return text.strip()
        except ImportError:
            # If markdown not installed, return raw
            return content.strip()
    except Exception as e:
        return f"Error reading Markdown: {str(e)}"


def extract_document_text(file_path: str) -> Dict[str, str]:
    """
    Extract text from any supported document type.

    Returns:
        Dict with 'text', 'filename', 'type', 'size', 'status'
    """
    path = Path(file_path)

    if not path.exists():
        return {
            'text': '',
            'filename': str(path.name),
            'type': 'unknown',
            'size': 0,
            'status': 'error',
            'error': 'File not found'
        }

    file_size = path.stat().st_size
    file_ext = path.suffix.lower()

    # Extract based on file type
    extractors = {
        '.pdf': extract_pdf_text,
        '.docx': extract_docx_text,
        '.doc': extract_docx_text,
        '.txt': extract_txt_text,
        '.md': extract_markdown_text,
        '.markdown': extract_markdown_text,
    }

    extractor = extractors.get(file_ext)

    if not extractor:
        return {
            'text': '',
            'filename': str(path.name),
            'type': file_ext,
            'size': file_size,
            'status': 'unsupported',
            'error': f'Unsupported file type: {file_ext}'
        }

    text = extractor(str(path))

    # Check if extraction was successful
    if text.startswith('Error'):
        return {
            'text': '',
            'filename': str(path.name),
            'type': file_ext,
            'size': file_size,
            'status': 'error',
            'error': text
        }

    return {
        'text': text,
        'filename': str(path.name),
        'type': file_ext,
        'size': file_size,
        'status': 'success',
        'char_count': len(text),
        'word_count': len(text.split())
    }


def extract_multiple_documents(file_paths: List[str]) -> Dict[str, any]:
    """
    Extract text from multiple documents.

    Returns:
        Dict with 'documents', 'combined_text', 'summary'
    """
    documents = []
    all_text = []
    total_chars = 0
    total_words = 0
    errors = []

    for file_path in file_paths:
        result = extract_document_text(file_path)
        documents.append(result)

        if result['status'] == 'success':
            all_text.append(f"\n\n=== {result['filename']} ===\n\n{result['text']}")
            total_chars += result['char_count']
            total_words += result['word_count']
        else:
            errors.append(f"{result['filename']}: {result.get('error', 'Unknown error')}")

    combined_text = "\n".join(all_text)

    return {
        'documents': documents,
        'combined_text': combined_text,
        'summary': {
            'total_files': len(file_paths),
            'successful': len([d for d in documents if d['status'] == 'success']),
            'failed': len([d for d in documents if d['status'] != 'success']),
            'total_chars': total_chars,
            'total_words': total_words,
            'errors': errors
        }
    }


def format_document_context(documents_result: Dict) -> str:
    """
    Format extracted documents for bot conversation context.

    Args:
        documents_result: Output from extract_multiple_documents()

    Returns:
        Formatted context string
    """
    summary = documents_result['summary']

    context = f"""
[DOCUMENT CONTEXT]
Loaded {summary['successful']}/{summary['total_files']} documents successfully
Total content: {summary['total_words']:,} words ({summary['total_chars']:,} characters)

Documents:
"""

    for doc in documents_result['documents']:
        if doc['status'] == 'success':
            context += f"  • {doc['filename']} ({doc['word_count']:,} words)\n"
        else:
            context += f"  • {doc['filename']} (failed: {doc.get('error', 'unknown')})\n"

    context += f"\n[DOCUMENT CONTENT]\n{documents_result['combined_text']}\n"

    return context


def get_supported_extensions() -> List[str]:
    """Get list of supported file extensions."""
    return ['.pdf', '.docx', '.doc', '.txt', '.md', '.markdown']


def is_supported_file(file_path: str) -> bool:
    """Check if file type is supported."""
    ext = Path(file_path).suffix.lower()
    return ext in get_supported_extensions()


if __name__ == "__main__":
    # Test module
    print("Document Processor Module")
    print("=" * 60)
    print(f"Supported formats: {', '.join(get_supported_extensions())}")
    print()

    # Example usage
    print("Usage:")
    print("  from document_processor import extract_document_text")
    print("  result = extract_document_text('document.pdf')")
    print("  print(result['text'])")
